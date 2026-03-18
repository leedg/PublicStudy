#!/usr/bin/env python3
"""
Build-LockSync-Report.py
동기화·락·비동기 처리 메커니즘 기술 분석 보고서 생성

사용법:
    python Build-LockSync-Report.py [출력경로.docx]

사전 요구사항:
    pip install python-docx
"""

from pathlib import Path
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── 색상 팔레트 ───────────────────────────────────────────────────────────────
C_TITLE_BG    = RGBColor(0x0D, 0x47, 0xA1)   # 진한 파랑 (타이틀 배경)
C_BLUE_DARK   = RGBColor(0x15, 0x65, 0xC0)   # 섹션 헤더
C_BLUE_LIGHT  = RGBColor(0xE3, 0xF2, 0xFD)   # note 박스 배경
C_INDIGO_DARK = RGBColor(0x1A, 0x23, 0x7E)   # 테이블 헤더 배경
C_PURPLE      = RGBColor(0x51, 0x2D, 0xA8)   # h3 색
C_ORANGE      = RGBColor(0xBF, 0x36, 0x0C)   # warning
C_YELLOW_BG   = RGBColor(0xFF, 0xF8, 0xE1)   # warning 배경
C_GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)   # tip
C_GREEN_BG    = RGBColor(0xE8, 0xF5, 0xE9)   # tip 배경
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK       = RGBColor(0x21, 0x21, 0x21)
C_GRAY        = RGBColor(0x75, 0x75, 0x75)
C_ROW_ALT     = RGBColor(0xF5, 0xF5, 0xF5)
C_CODE_BG     = RGBColor(0xF5, 0xF5, 0xF5)
C_CODE_FG     = RGBColor(0x37, 0x47, 0x4F)
C_CAPTION     = RGBColor(0x61, 0x61, 0x61)
C_TITLE_STRIP = RGBColor(0x1E, 0x88, 0xE5)   # 타이틀 악센트 선

FONT      = "맑은 고딕"
FONT_CODE = "Consolas"
ASSETS    = Path(__file__).parent.parent / "assets"


# ── XML 유틸리티 ──────────────────────────────────────────────────────────────
def _rgb_hex(c: RGBColor) -> str:
    # RGBColor is a subclass of tuple: (r, g, b)
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def shade_cell(cell, color: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    tcPr.append(shd)


def shade_para(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    pPr.append(shd)


def add_left_border(para, color_hex: str = "1565C0", sz: int = 12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def add_bottom_border(para, color_hex: str = "CCCCCC", sz: int = 4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_cell_valign_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


# ── Run/폰트 헬퍼 ──────────────────────────────────────────────────────────────
def _apply_font(run, name=FONT, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass
    return run


def add_run(para, text, name=FONT, size=None, bold=False, italic=False, color=None):
    return _apply_font(para.add_run(text), name, size, bold, italic, color)


# ── 문서 셋업 ─────────────────────────────────────────────────────────────────
def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    style_cfg = [
        ("Normal",       10.5),
        ("Heading 1",    15),
        ("Heading 2",    13),
        ("Heading 3",    11.5),
        ("List Bullet",  10.5),
        ("List Number",  10.5),
    ]
    for name, sz in style_cfg:
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(sz)
        try:
            s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        except Exception:
            pass
    return doc


# ── 타이틀 페이지 ─────────────────────────────────────────────────────────────
def build_title_page(doc: Document):
    def _shade(text="", align=WD_ALIGN_PARAGRAPH.LEFT,
               before=0, after=0, indent=0,
               bold=False, italic=False, fsize=None, fcolor=None, bg=C_TITLE_BG):
        p = doc.add_paragraph(style="Normal")
        shade_para(p, bg)
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if text:
            add_run(p, text, bold=bold, italic=italic, size=fsize, color=fcolor)
        return p

    # 상단 여백
    for _ in range(3):
        _shade()

    # 악센트 가로선
    accent = _shade(before=2, after=2, bg=C_TITLE_STRIP)
    p_bdr = accent._p.get_or_add_pPr()
    shd_b = OxmlElement("w:shd")
    shd_b.set(qn("w:val"), "clear"); shd_b.set(qn("w:color"), "auto")
    shd_b.set(qn("w:fill"), _rgb_hex(C_TITLE_STRIP))
    p_bdr.append(shd_b)

    # 메인 타이틀
    _shade("동기화·락·비동기 처리",
           align=WD_ALIGN_PARAGRAPH.CENTER, before=30, after=4,
           bold=True, fsize=28, fcolor=C_WHITE)
    _shade("메커니즘 분석 보고서",
           align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=20,
           bold=True, fsize=28, fcolor=C_WHITE)

    # 부제
    _shade("NetworkModuleTest — ServerEngine 구현 기반",
           align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=30,
           italic=True, fsize=14,
           fcolor=RGBColor(0xBB, 0xDE, 0xFB))

    # 구분선 (얇은 밝은 파랑)
    _shade(before=0, after=16, bg=RGBColor(0x42, 0x86, 0xC5))

    # 메타 정보
    meta = [
        ("기준 리포지토리", "NetworkModuleTest"),
        ("분석 대상",      "ServerEngine — Concurrency / Network / Utils"),
        ("작성일",         "2026-03-10"),
        ("버전",           "1.0"),
    ]
    for label, value in meta:
        p = _shade(before=0, after=4, indent=4)
        add_run(p, f"{label}:  ", bold=True, size=11,
                color=RGBColor(0x90, 0xCA, 0xF9))
        add_run(p, value, size=11, color=C_WHITE)

    # 하단 여백
    for _ in range(5):
        _shade()

    doc.add_page_break()


# ── 본문 빌딩 블록 ────────────────────────────────────────────────────────────
def h1(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    shade_para(p, RGBColor(0xE8, 0xEA, 0xF6))
    add_left_border(p, color_hex="0D47A1", sz=20)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    add_run(p, text, bold=True, size=15, color=C_TITLE_BG)


def h2(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    add_left_border(p, color_hex="1565C0", sz=12)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, size=13, color=C_BLUE_DARK)


def h3(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text, bold=True, size=11.5, color=C_PURPLE)


def body(doc: Document, text: str, size: float = 10.5,
         color: RGBColor = C_BLACK, indent_cm: float = 0) -> object:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size=size, color=color)
    return p


def bullet(doc: Document, text: str, size: float = 10.5, indent_cm: float = 0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size=size)
    return p


def code_ref(doc: Document, text: str, note: str = ""):
    p = doc.add_paragraph(style="Normal")
    shade_para(p, C_CODE_BG)
    add_left_border(p, color_hex="90A4AE", sz=6)
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, name=FONT_CODE, size=9.5, bold=True, color=C_CODE_FG)
    if note:
        add_run(p, f"  — {note}", size=9, color=C_GRAY)
    return p


def callout(doc: Document, text: str, style: str = "note"):
    cfg = {
        "note":    (C_BLUE_LIGHT, C_BLUE_DARK,  "ℹ  참고"),
        "warning": (C_YELLOW_BG,  C_ORANGE,     "⚠  주의"),
        "tip":     (C_GREEN_BG,   C_GREEN_DARK, "✔  핵심"),
    }
    bg, fg, label = cfg.get(style, cfg["note"])
    p = doc.add_paragraph(style="Normal")
    shade_para(p, bg)
    add_left_border(p, color_hex=_rgb_hex(fg), sz=8)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, f"{label}:  ", bold=True, size=10, color=fg)
    add_run(p, text, size=10, color=C_BLACK)
    return p


def image(doc: Document, filename: str, caption: str = None,
          width: float = 5.8):
    path = ASSETS / filename
    if not path.exists():
        body(doc, f"[이미지 없음: {filename}]", color=RGBColor(0xC6, 0x28, 0x28))
        return
    doc.add_picture(str(path), width=Inches(width))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap_p = doc.add_paragraph(style="Normal")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(2)
        cap_p.paragraph_format.space_after  = Pt(10)
        add_run(cap_p, f"▲  {caption}", italic=True, size=9, color=C_CAPTION)


def make_table(doc: Document, headers: list, rows: list,
               col_widths_cm: list = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, C_INDIGO_DARK)
        set_cell_valign_center(cell)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, color=C_WHITE, size=9.5)

    # 데이터 행
    for ri, row_data in enumerate(rows, 1):
        bg = C_ROW_ALT if ri % 2 == 0 else None
        for ci, val in enumerate(row_data):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            if bg:
                shade_cell(cell, bg)
            set_cell_valign_center(cell)
            p = cell.paragraphs[0]
            add_run(p, str(val), size=9.5)

    # 열 너비
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t


def divider(doc: Document):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    add_bottom_border(p, color_hex="CCCCCC", sz=4)


# ── 섹션별 본문 ───────────────────────────────────────────────────────────────
def section_1(doc: Document):
    h1(doc, "1. 개요")
    body(doc, (
        "ServerEngine은 Windows · Linux · macOS를 동시 지원하는 다중 플랫폼 네트워크 엔진이다. "
        "수천 개의 동시 세션을 처리하면서 데이터 경합과 데드락 없이 안전하게 동작하기 위해 "
        "아래 세 가지 설계 원칙을 따른다."
    ))
    bullet(doc, "역할 분리:  I/O 완료 스레드는 수신·송신만 처리, 패킷 로직 실행은 KeyedDispatcher 워커 스레드에서 전담")
    bullet(doc, "락 중첩 없음:  모든 뮤텍스는 단일 책임으로 획득되며, 뮤텍스를 보유한 채 다른 뮤텍스를 획득하는 코드가 없다")
    bullet(doc, "락 최소화:  핫 패스는 atomic CAS / lock-free 큐를 우선 사용하고, 뮤텍스는 cold path로 제한")
    callout(doc, (
        "LockProfiling 빌드 플래그(NET_LOCK_PROFILING)를 활성화하면 "
        "모든 뮤텍스의 대기 시간(waitNs)과 보유 시간(holdNs)을 런타임에 수집할 수 있다. "
        "비활성화 시 표준 std::lock_guard / std::unique_lock으로 zero-overhead 컴파일된다."
    ), style="tip")
    divider(doc)


def section_2(doc: Document):
    h1(doc, "2. 동기화 프리미티브 총람")
    body(doc, (
        "ServerEngine 코드베이스에서 사용하는 모든 동기화 프리미티브를 소유 클래스별로 정리한다. "
        "아래 다이어그램은 각 클래스가 어떤 뮤텍스·atomic 변수를 소유하는지 한눈에 보여준다."
    ))
    image(doc, "07-lock-structure.png",
          caption="동기화 프리미티브 소유 관계 — 클래스별 Mutex 및 Atomic 변수 전체 맵")

    h2(doc, "2.1 Mutex 사용 현황")
    make_table(doc,
        ["뮤텍스", "소유 클래스", "보호 대상", "설계 포인트"],
        [
            ["mSendMutex",       "Session",             "mSendQueue\nmAsyncProvider",
             "lock 내 shared_ptr 스냅샷만 복사 후 즉시 해제\n실제 I/O는 락 외부에서 호출"],
            ["mCallbackMutex",   "BaseNetworkEngine",   "mCallbacks 이벤트 맵",
             "이벤트 등록·해제 시에만 짧게 획득"],
            ["mDrainMutex",      "AsyncScope",          "WaitForDrain 조건 대기",
             "notify_all은 뮤텍스 없이 호출\n(WaitForDrain과 불필요한 경합 방지)"],
            ["mMutex",           "TimerQueue",          "mHeap · mCancelledHandles",
             "min-heap 수정 + 취소 핸들 집합 일관 관리"],
            ["mMutexQueueMutex", "ExecutionQueue",      "mMutexQueue (Mutex 백엔드)", ""],
            ["mWaitMutex",       "ExecutionQueue",      "mNotEmptyCV · mNotFullLFCV",
             "Mutex / LockFree 백엔드 공통\n대기 CV 보호"],
        ],
        col_widths_cm=[3.8, 4.4, 4.2, 5.6]
    )
    callout(doc, (
        "mWorkersMutex (KeyedDispatcher)는 std::shared_mutex다. "
        "Dispatch()는 shared lock을 획득해 다수의 I/O 워커가 동시에 디스패치할 수 있고, "
        "Shutdown()은 exclusive lock을 획득해 mWorkers.clear() 중 Dispatch()가 벡터에 접근하지 못하게 한다."
    ), style="note")

    h2(doc, "2.2 Atomic 변수 및 Memory Ordering")
    body(doc, "핫 패스에서 뮤텍스 대신 atomic을 사용하는 변수와 memory_order 선택 근거다.")
    make_table(doc,
        ["Atomic 변수", "소유 클래스", "역할", "사용 순서 및 이유"],
        [
            ["mState",                  "Session",              "연결 상태 (enum)",
             "acq_rel exchange\nClose()의 TOCTOU 이중 닫기 방지"],
            ["mSocket",                 "Session",              "소켓 핸들",
             "acq_rel exchange\nClose/Send 간 race 방지"],
            ["mIsSending",              "Session",              "CAS 이중전송 방지",
             "compare_exchange_strong\n한 스레드만 PostSend() 진입 보장"],
            ["mSendQueueSize",          "Session",              "lock-free 큐 크기 조회",
             "relaxed read (fast-path)\nrelease store"],
            ["mPingSequence",           "Session",              "핑 시퀀스 카운터",
             "relaxed\n순서 보장 불필요, 단순 카운터"],
            ["mCancelled",              "AsyncScope",           "협력 취소 플래그",
             "release store / acquire load"],
            ["mInFlight",               "AsyncScope",           "in-flight 작업 수",
             "acq_rel fetch_add/sub\nEndTask 완료 감지 및 notify"],
            ["mEnqueuePos / mDequeuePos", "BoundedLockFreeQueue", "MPMC 링 버퍼 위치",
             "relaxed CAS\n+ acquire seq load\n+ release seq store"],
            ["mRunning",                "KeyedDispatcher",      "실행 상태 플래그",
             "acq_rel compare_exchange_strong"],
        ],
        col_widths_cm=[4.2, 4.0, 3.8, 6.0]
    )

    h3(doc, "Memory Order 빠른 참조")
    make_table(doc,
        ["memory_order", "의미", "주요 사용 패턴"],
        [
            ["relaxed", "원자성만 보장, 재배치 허용", "mPingSequence, 통계 카운터"],
            ["acquire", "이 load 이후 접근이 앞으로 이동 불가", "mState/mSocket load, IsCancelled()"],
            ["release", "이 store 이전 쓰기가 뒤로 이동 불가", "mState/mSocket store, BoundedLockFreeQueue 시퀀스 store"],
            ["acq_rel", "acquire + release 동시 (RMW 전용)", "exchange, fetch_sub, compare_exchange (mState, mInFlight 등)"],
        ],
        col_widths_cm=[2.8, 6.2, 9.0]
    )
    divider(doc)


def section_3(doc: Document):
    h1(doc, "3. AsyncScope — 협력 취소와 생명주기 관리")
    body(doc, (
        "Session은 AsyncScope를 하나씩 멤버로 갖는다. "
        "로직 스레드풀에 디스패치된 람다가 Session 소멸 후에도 실행되는 use-after-free를 방지하는 것이 핵심 목적이다. "
        "RAII 소멸자가 Cancel() + WaitForDrain()을 자동 호출하므로 별도 정리 코드가 불필요하다."
    ))

    h2(doc, "3.1 설계 목적")
    bullet(doc, "BaseNetworkEngine은 Dispatch() 대신 mAsyncScope.Submit()을 호출한다")
    bullet(doc, "Submit은 IsCancelled()를 확인 후 람다를 큐잉한다 — Cancelled 상태이면 람다 실행 없이 in-flight를 즉시 감소")
    bullet(doc, "RAII 소멸자가 Cancel() + WaitForDrain()을 자동 호출하여 Session 소멸 시 안전하게 정리된다")

    h2(doc, "3.2 상태 전이 다이어그램")
    image(doc, "08-async-scope-lifecycle.png",
          caption="AsyncScope 생명주기 상태 전이 — 생성 / Cancelled / Drained / Reset(풀 재사용)")

    h2(doc, "3.3 상태별 동작 요약")
    make_table(doc,
        ["상태", "진입 조건", "Submit() 동작", "특이사항"],
        [
            ["Active",    "생성 직후 / Reset() 호출",
             "in-flight++ 후 람다 큐잉 (정상)",
             "mCancelled = false"],
            ["Cancelled", "Cancel() 호출",
             "fast-path: in-flight++ 후 즉시 --\n(큐잉 없이 skip)",
             "mCancelled = true (release 순서)"],
            ["Drained",   "WaitForDrain() 반환",
             "Submit 불가 (이미 Cancelled)",
             "mInFlight == 0 보장됨"],
        ],
        col_widths_cm=[2.5, 4.0, 5.8, 5.7]
    )

    h2(doc, "3.4 풀 재사용 시 주의사항")
    callout(doc, (
        "Session::Reset() 호출 시 반드시 mAsyncScope.Reset()을 함께 호출해야 한다. "
        "누락 시 재사용 세션의 mCancelled=true가 잔존하여 모든 수신 람다가 silently skip된다. "
        "서버가 응답 불가 상태에 빠지면서 클라이언트는 EAGAIN을 반복 수신하게 된다. "
        "(이 버그는 io_uring 백엔드에서 실제로 발생한 이력이 있다 — 2026-03-02 수정)"
    ), style="warning")
    body(doc, "Reset() 선행 조건: mInFlight == 0. SessionPool::ReleaseInternal이 마지막 shared_ptr 참조를 해제하기 직전에 호출하므로 이 조건은 구조적으로 보장된다.")
    code_ref(doc, "AsyncScope.h:46  →  void Reset() { assert(mInFlight==0); mCancelled.store(false, release); }")
    divider(doc)


def section_4(doc: Document):
    h1(doc, "4. KeyedDispatcher — 키 친화도 라우팅")
    body(doc, (
        "KeyedDispatcher는 sessionId를 key로 사용해 동일 세션의 모든 로직 작업을 "
        "항상 같은 워커 큐로 라우팅한다. 워커 큐의 FIFO 특성과 결합하여 세션 단위 처리 순서를 보장한다."
    ))

    h2(doc, "4.1 라우팅 다이어그램")
    image(doc, "09-keyed-dispatcher-routing.png",
          caption="KeyedDispatcher 키 친화도 라우팅 — I/O 완료 이벤트 → 로직 워커 분배 흐름")

    h2(doc, "4.2 핵심 설계 포인트")
    make_table(doc,
        ["항목", "내용"],
        [
            ["라우팅 공식",       "workerIndex = sessionId % workerCount"],
            ["Dispatch 락",      "shared_lock(mWorkersMutex)\n다수의 I/O 워커가 동시 디스패치 허용"],
            ["Shutdown 락",      "exclusive lock(mWorkersMutex)\nmWorkers.clear() 중 Dispatch() 차단"],
            ["TOCTOU 방어",      "shared lock 범위 내에서 mRunning 확인 + mWorkers 접근\n→ clear()와 경쟁 없음"],
            ["mRecvMutex 제거",  "동일 세션 → 동일 워커 → ProcessRawRecv 순차 실행 보장\n→ 수신 누적 버퍼 락 불필요"],
            ["mRecvBatchBuf",    "KeyedDispatcher 친화도가 동시 접근을 구조적으로 배제\n→ 락 없이 안전한 배치 버퍼 재사용"],
        ],
        col_widths_cm=[4.0, 14.0]
    )
    callout(doc, (
        "workerCount=1로 설정하면 모든 세션을 단일 워커로 직렬화할 수 있다. "
        "DBTaskQueue가 Initialize(1)로 워커 1개를 강제하는 이유가 여기에 있다 — "
        "같은 sessionId에 대한 DB 작업의 순서를 보장한다."
    ), style="tip")
    divider(doc)


def section_5(doc: Document):
    h1(doc, "5. ExecutionQueue — 이중 백엔드")
    body(doc, (
        "ExecutionQueue<T>는 KeyedDispatcher의 워커 큐로 사용된다. "
        "QueueBackend(Mutex / LockFree)와 BackpressurePolicy(RejectNewest / Block)의 "
        "조합으로 4가지 동작 모드를 지원한다."
    ))

    h2(doc, "5.1 아키텍처 다이어그램")
    image(doc, "10-execution-queue-backends.png",
          caption="ExecutionQueue 이중 백엔드 구조 — Mutex / Lock-Free 백엔드 및 CV 대기 흐름")

    h2(doc, "5.2 백엔드 비교")
    make_table(doc,
        ["항목", "Mutex 백엔드", "Lock-Free 백엔드"],
        [
            ["자료구조",          "std::queue<T>\n+ mMutexQueueMutex",
             "BoundedLockFreeQueue<T>\n링 버퍼 (2의 거듭제곱 크기)"],
            ["크기 제한",         "선택적 (0 = 무제한 가능)",
             "필수 (기본 1024, 2^N으로 보정)"],
            ["스레드 모델",       "MPMC (mutex 직렬화)",
             "MPMC (CAS 기반, 무락 enqueue/dequeue)"],
            ["Block 정책",        "mNotFullMutexCV로 생산자 대기",
             "spin + mNotFullLFCV로 대기"],
            ["mSize 정확도",      "정확 (mutex 보호)",
             "±1 근사 — fetch_add와 TryEnqueue가 비원자\n모니터링 전용으로만 사용"],
            ["적합한 용도",       "일반 작업 큐, 크기 유연성 필요",
             "지연 최소화가 중요한 고빈도 작업"],
        ],
        col_widths_cm=[3.5, 7.0, 7.5]
    )

    h2(doc, "5.3 BoundedLockFreeQueue 링 버퍼 원리")
    body(doc, (
        "각 Cell은 sequence atomic을 갖는다. enqueuePos / dequeuePos와 sequence의 차이로 슬롯 상태를 판별한다. "
        "정상 CAS 성공 → 슬롯 선점 → 데이터 쓰기/읽기 → sequence 갱신(release) 순서로 동작한다."
    ))
    make_table(doc,
        ["조건 (enqueue 시)", "의미", "조건 (dequeue 시)", "의미"],
        [
            ["seq − pos == 0",  "슬롯 비어있음\n→ CAS로 선점",
             "seq − (pos+1) == 0", "데이터 있음\n→ CAS로 선점"],
            ["seq − pos  < 0",  "큐 가득 참\n→ false 반환",
             "seq − (pos+1) < 0",  "큐 비어있음\n→ false 반환"],
            ["seq − pos  > 0",  "경쟁 패배\n→ pos 재로드 후 재시도",
             "seq − (pos+1) > 0",  "경쟁 패배\n→ pos 재로드 후 재시도"],
        ],
        col_widths_cm=[3.5, 4.2, 3.8, 6.5]
    )
    callout(doc, (
        "Cell은 alignas(64)로 캐시라인(64 byte) 정렬된다. "
        "인접 슬롯 간 false sharing을 방지하여 MPMC 경합 시 CPU 캐시 무효화로 인한 성능 저하를 막는다."
    ), style="tip")
    divider(doc)


def section_6(doc: Document):
    h1(doc, "6. 통합 흐름 — 수신부터 로직 처리까지")
    body(doc, (
        "수신 패킷 한 개가 처리되는 전체 경로에서 각 동기화 프리미티브가 언제·어떻게 사용되는지 단계별로 정리한다."
    ))
    make_table(doc,
        ["단계", "실행 스레드", "동기화 수단", "코드 위치"],
        [
            ["recv 완료 감지",
             "I/O 워커\n(IOCP/epoll/io_uring)",
             "없음 (플랫폼 I/O 완료)",
             "ProcessCompletions()"],
            ["ProcessRawRecv 디스패치",
             "I/O 워커",
             "mAsyncScope.Submit()\nin-flight++ + IsCancelled 확인",
             "BaseNetworkEngine.cpp:255"],
            ["PacketHeader 파싱",
             "KeyedDispatcher 워커",
             "없음\n(동일 세션 → 동일 워커 구조 보장)",
             "Session.cpp:535"],
            ["OnRecv 콜백 실행",
             "KeyedDispatcher 워커",
             "없음",
             "Session.cpp:563"],
            ["Session::Send() 호출",
             "임의 스레드 (로직 워커 등)",
             "mSendQueueSize relaxed read\n→ mSendMutex lock\n→ CAS mIsSending",
             "Session.cpp:199"],
            ["PostSend() / WSASend",
             "CAS 선점 스레드",
             "CAS(mIsSending)\n→ mSendMutex lock\n→ SendBufferPool slot acquire",
             "Session.cpp:317"],
            ["송신 완료 / 플래그 해제",
             "I/O 워커",
             "mIsSending.store(false, release)\n+ TOCTOU 재확인(큐 재확인 후 재시도)",
             "Session.cpp:336"],
        ],
        col_widths_cm=[3.8, 3.8, 5.6, 4.8]
    )
    callout(doc, (
        "Close() 경로: Session::Close()는 mState.exchange(Disconnected, acq_rel)로 이중 닫기를 방지하고, "
        "mSendMutex 단일 락으로 mAsyncProvider 해제와 mSendQueue 드레인을 원자적으로 수행한다. "
        "마지막으로 mAsyncScope.Cancel()로 큐잉 중인 람다를 협력 취소한다."
    ), style="note")
    divider(doc)


def section_7(doc: Document):
    h1(doc, "7. 주요 코드 참조")
    make_table(doc,
        ["파일 (Server/ServerEngine/...)", "관련 섹션"],
        [
            ["Utils/LockProfiling.h",              "2.1 — LockProfiling 매크로 (NET_LOCK_PROFILING)"],
            ["Network/Core/Session.h",             "2.1, 2.2, 6 — mSendMutex · CAS 이중전송"],
            ["Network/Core/Session.cpp",           "6 — 통합 흐름 전체"],
            ["Concurrency/AsyncScope.h",           "3 — AsyncScope 전체"],
            ["Concurrency/KeyedDispatcher.h",      "4 — 라우팅 · shared_mutex"],
            ["Concurrency/ExecutionQueue.h",       "5 — 이중 백엔드"],
            ["Concurrency/BoundedLockFreeQueue.h", "5.3 — 링 버퍼 원리"],
            ["Concurrency/TimerQueue.h/.cpp",      "2.1 — mMutex · min-heap"],
            ["Network/Core/BaseNetworkEngine.h",   "2.1, 4 — mCallbackMutex · mLogicDispatcher"],
            ["Utils/SafeQueue.h",                  "ThreadPool 기반 작업 큐"],
            ["Utils/ThreadPool.h",                 "SafeQueue 기반 스레드 풀"],
        ],
        col_widths_cm=[9.5, 8.5]
    )


# ── main ──────────────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    else:
        out = Path(__file__).parent.parent / "Lock_Sync_Async_Mechanisms.docx"

    doc = setup_document()
    build_title_page(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
