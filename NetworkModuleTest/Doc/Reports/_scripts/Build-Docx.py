# ==============================================================================
# Build-Docx.py  (Doc/Reports/_scripts/Build-Docx.py)
# 역할: 마크다운 파일(.md)을 Word 문서(.docx)로 변환한다.
#       Doc/Reports 패키지(ExecutiveSummary, TeamShare, WikiPackage)의
#       DOCX 파일을 생성하는 데 사용한다.
#
# 기능:
#   - 헤딩(#, ##, ###), 리스트(-, 1.), 코드블록(```), 표(|), 이미지(![...]) 지원
#   - 이미지는 .svg → .png 자동 전환 (같은 경로에 .png 있을 경우)
#   - 한글 폰트(맑은 고딕) 및 각 스타일별 글자 크기 자동 설정
#   - 단일 파일 → 단일 docx (single 모드)
#   - 여러 파일 → 페이지 구분으로 묶인 단일 docx (bundle 모드)
#
# 사용법:
#   python Build-Docx.py single <입력.md> <출력.docx>
#   python Build-Docx.py bundle <제목> <출력.docx> <입력1.md> [입력2.md ...]
#
# 사전 요구사항:
#   pip install python-docx
# ==============================================================================
from pathlib import Path
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


FONT_NAME = "Malgun Gothic"


def apply_run_font(run, font_name: str = FONT_NAME, size: Pt | None = None):
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_style(style, font_name: str = FONT_NAME, size: Pt | None = None):
    style.font.name = font_name
    if size is not None:
        style.font.size = size
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure(doc: Document):
    for style_name, font_size in [
        ("Normal", Pt(10.5)),
        ("Title", Pt(20)),
        ("Heading 1", Pt(16)),
        ("Heading 2", Pt(13)),
        ("Heading 3", Pt(11.5)),
        ("List Bullet", Pt(10.5)),
        ("List Number", Pt(10.5)),
    ]:
        configure_style(doc.styles[style_name], size=font_size)


def add_paragraph_with_inline(doc: Document, text: str):
    p = doc.add_paragraph()
    apply_run_font(p.add_run(text))


def resolve_image(md_file: Path, rel: str) -> Path | None:
    candidate = (md_file.parent / rel).resolve()
    if candidate.suffix.lower() == ".svg":
        png = candidate.with_suffix(".png")
        if png.exists():
            return png
    if candidate.exists():
        return candidate
    return None


def parse_table(doc: Document, lines: list[str]):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) < 2:
        for line in lines:
            doc.add_paragraph(line)
        return

    header = rows[0]
    body = rows[2:] if len(rows) >= 3 else []
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for i, value in enumerate(header):
        cell = table.cell(0, i)
        cell.text = ""
        apply_run_font(cell.paragraphs[0].add_run(value))
    for r, row in enumerate(body, start=1):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            apply_run_font(cell.paragraphs[0].add_run(value))


def render_markdown(doc: Document, md_file: Path):
    lines = md_file.read_text(encoding="utf-8").splitlines()
    in_code = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            apply_run_font(run, font_name="Consolas", size=Pt(9))
            i += 1
            continue

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            rel = stripped.split("](", 1)[1][:-1]
            img = resolve_image(md_file, rel)
            if img is not None:
                doc.add_picture(str(img), width=Inches(6.3))
            else:
                doc.add_paragraph(stripped)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            parse_table(doc, block)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped[:3] == "1. ":
            doc.add_paragraph(stripped[3:].strip(), style="List Number")
        else:
            add_paragraph_with_inline(doc, line)
        i += 1


def build_single(source: Path, out_path: Path):
    doc = Document()
    configure(doc)
    render_markdown(doc, source)
    doc.save(out_path)


def build_bundle(sources: list[Path], out_path: Path, title: str | None = None):
    doc = Document()
    configure(doc)
    if title:
        doc.add_heading(title, level=1)
    for idx, source in enumerate(sources):
        if idx > 0:
            doc.add_page_break()
        render_markdown(doc, source)
    doc.save(out_path)


if __name__ == "__main__":
    mode = sys.argv[1]
    if mode == "single":
        build_single(Path(sys.argv[2]), Path(sys.argv[3]))
    elif mode == "bundle":
        title = sys.argv[2]
        out_path = Path(sys.argv[3])
        sources = [Path(x) for x in sys.argv[4:]]
        build_bundle(sources, out_path, title)
    else:
        raise SystemExit("unknown mode")
