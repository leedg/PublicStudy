#!/usr/bin/env python3
"""
Build-NetworkAsyncDB-Report.py
네트워크 / 비동기 처리 / DB 처리 구조 분석 보고서 생성

사용법:
    python Build-NetworkAsyncDB-Report.py [출력경로.docx]

사전 요구사항:
    pip install python-docx

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
diagrams/ drawio 원본  →  assets/ PNG (보고서에 삽입되는 파일)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Sec01-Overview-ArchitectureLayers.drawio       → diag_arch.png
  └ PIL 재작성(_scripts/Fix-ReviewCorrections.py draw_arch)
Sec01-Overview-SystemFlowChart.drawio          → diag_overview.png
  └ draw.io 수동 내보내기
Sec02-Network-ClientLifecycleSequence.drawio   → diag_client_lifecycle.png
  └ draw.io 수동 내보내기
Sec02-Network-ConnectionEstablishmentSequence.drawio → diag_seq.png
  └ SVG+Chrome 렌더(Fix-ReviewCorrections.py render_seq)
Sec02-Network-SessionUML.drawio                → diag_session_uml.png
  └ draw.io 수동 내보내기
Sec03-Async-IOCompletionDispatchFlow.drawio    → diag_async_1_dispatch.png
  └ PIL(Fix-ReviewCorrections.py draw_dispatch), H=600
Sec03.3-AsyncWAL-KeyedDispatcherSessionRouting.drawio → diag_async_2_keyed.png
  └ PIL(Fix-ReviewCorrections.py draw_keyed), H=620
Sec04-DB-ProcessingLayerWALFactoryQueue.drawio → diag_db.png
  └ draw.io 수동 내보내기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PNG 재생성:  python _scripts/Fix-ReviewCorrections.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

from pathlib import Path
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── 색상 팔레트 ───────────────────────────────────────────────────────────────
C_TITLE_BG    = RGBColor(0x0D, 0x47, 0xA1)   # 진한 파랑 (타이틀 배경)
C_BLUE_DARK   = RGBColor(0x15, 0x65, 0xC0)   # 섹션 헤더
C_BLUE_LIGHT  = RGBColor(0xE3, 0xF2, 0xFD)   # note 박스 배경
C_TEAL_DARK   = RGBColor(0x00, 0x69, 0x5C)   # h1 네트워크 섹션 강조
C_INDIGO_DARK = RGBColor(0x1A, 0x23, 0x7E)   # 테이블 헤더 배경
C_PURPLE      = RGBColor(0x51, 0x2D, 0xA8)   # h3 색
C_ORANGE      = RGBColor(0xBF, 0x36, 0x0C)   # warning
C_YELLOW_BG   = RGBColor(0xFF, 0xF8, 0xE1)   # warning 배경
C_GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)   # tip
C_GREEN_BG    = RGBColor(0xE8, 0xF5, 0xE9)   # tip 배경
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK       = RGBColor(0x21, 0x21, 0x21)
C_GRAY        = RGBColor(0x75, 0x75, 0x75)
C_ROW_ALT     = RGBColor(0xF5, 0xF5, 0xF5)
C_CODE_BG     = RGBColor(0xF5, 0xF5, 0xF5)
C_CODE_FG     = RGBColor(0x37, 0x47, 0x4F)
C_CAPTION     = RGBColor(0x61, 0x61, 0x61)
C_TITLE_STRIP = RGBColor(0x1E, 0x88, 0xE5)   # 타이틀 악센트 선

FONT      = "맑은 고딕"
FONT_CODE = "Consolas"
# draw.io 다이어그램 스타일 — "orig" | "B" | "C"
# main()에서 CLI 인수로 오버라이드 가능: python Build-NetworkAsyncDB-Report.py [출력.docx] [orig|B|C]
DRAWIO_STYLE = "orig"
# 기존 다이어그램 (ExecutiveSummary/assets/)  — Reports/../ExecutiveSummary/assets
ASSETS_EXEC = Path(__file__).parent.parent.parent / "ExecutiveSummary" / "assets"
# draw.io 내보낸 다이어그램 원본 (assets/)    — NetworkAsyncDBReport/assets
ASSETS_NEW  = Path(__file__).parent.parent / "assets"


# ── XML 유틸리티 ──────────────────────────────────────────────────────────────
def _rgb_hex(c: RGBColor) -> str:
    # RGBColor is a subclass of tuple: (r, g, b)
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"


def shade_cell(cell, color: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    tcPr.append(shd)


def shade_para(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    pPr.append(shd)


def add_left_border(para, color_hex: str = "1565C0", sz: int = 12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def add_bottom_border(para, color_hex: str = "CCCCCC", sz: int = 4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_cell_valign_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


# ── Run/폰트 헬퍼 ──────────────────────────────────────────────────────────────
def _apply_font(run, name=FONT, size=None, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass
    return run


def add_run(para, text, name=FONT, size=None, bold=False, italic=False, color=None):
    return _apply_font(para.add_run(text), name, size, bold, italic, color)


# ── 문서 셋업 ─────────────────────────────────────────────────────────────────
def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    style_cfg = [
        ("Normal",      10.5, None,       False),
        ("List Bullet", 10.5, None,       False),
        ("List Number", 10.5, None,       False),
    ]
    for name, sz, color, bold in style_cfg:
        s = doc.styles[name]
        s.font.name = FONT
        s.font.size = Pt(sz)
        s.font.bold = bold
        if color:
            s.font.color.rgb = color
        try:
            s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        except Exception:
            pass

    # ── Heading 스타일: Word TOC 자동 인식용 ─────────────────────────────
    # h1/h2/h3 함수가 이 스타일을 사용하므로 Word의 자동 목차가 작동한다.
    heading_cfg = [
        ("Heading 1", 15,   C_TITLE_BG,  True),
        ("Heading 2", 13,   C_BLUE_DARK, True),
        ("Heading 3", 11.5, C_PURPLE,    True),
    ]
    for name, sz, color, bold in heading_cfg:
        s = doc.styles[name]
        s.font.name  = FONT
        s.font.size  = Pt(sz)
        s.font.bold  = bold
        s.font.color.rgb = color
        # 번호 매기기 제거 (Word 기본 Heading 스타일은 번호 있음)
        try:
            pPr = s._element.get_or_add_pPr()
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)
        except Exception:
            pass
        try:
            s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        except Exception:
            pass

    # ── 문서 열기 시 필드 자동 업데이트 (TOC 갱신 포함) ─────────────────────
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)

    return doc


# ── 타이틀 페이지 ─────────────────────────────────────────────────────────────
def build_title_page(doc: Document):
    def _blank(n: int = 1):
        for _ in range(n):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

    # ── 상단 여백 ─────────────────────────────────────────────────────────
    _blank(5)

    # ── 메인 타이틀 (색상 없음, 굵은 텍스트만) ────────────────────────────
    for line in ("네트워크, 비동기처리, DB처리",):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        add_run(p, line, bold=True, size=28, color=C_BLACK)

    # ── 중간 여백 ─────────────────────────────────────────────────────────
    _blank(18)

    # ── 이름 (우측 하단) ──────────────────────────────────────────────────
    p_name = doc.add_paragraph(style="Normal")
    p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after  = Pt(0)
    add_run(p_name, "이름 : 이동길", bold=False, size=12, color=C_BLACK)

    doc.add_page_break()


# ── 본문 빌딩 블록 ────────────────────────────────────────────────────────────
def h1(doc: Document, text: str):
    # Word "Heading 1" 스타일 사용 → 자동 목차(TOC) 인식
    p = doc.add_paragraph(style="Heading 1")
    shade_para(p, RGBColor(0xE8, 0xEA, 0xF6))
    add_left_border(p, color_hex="0D47A1", sz=20)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 스타일에서 색상/폰트를 이미 설정했으므로 run의 명시적 override 없이 텍스트만 추가
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = C_TITLE_BG
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except Exception:
        pass


def h2(doc: Document, text: str):
    # Word "Heading 2" 스타일 사용 → 자동 목차(TOC) 인식
    p = doc.add_paragraph(style="Heading 2")
    add_left_border(p, color_hex="1565C0", sz=12)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_BLUE_DARK
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except Exception:
        pass


def h3(doc: Document, text: str):
    # Word "Heading 3" 스타일 사용 → 자동 목차(TOC) 인식
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = C_PURPLE
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except Exception:
        pass


def body(doc: Document, text: str, size: float = 10.5,
         color: RGBColor = C_BLACK, indent_cm: float = 0) -> object:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size=size, color=color)
    return p


def bullet(doc: Document, text: str, size: float = 10.5, indent_cm: float = 0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size=size)
    return p


def numbered(doc: Document, text: str, size: float = 10.5, indent_cm: float = 0.5):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    add_run(p, text, size=size)
    return p


def code_ref(doc: Document, text: str, note: str = ""):
    p = doc.add_paragraph(style="Normal")
    shade_para(p, C_CODE_BG)
    add_left_border(p, color_hex="90A4AE", sz=6)
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, name=FONT_CODE, size=9.5, bold=True, color=C_CODE_FG)
    if note:
        add_run(p, f"  — {note}", size=9, color=C_GRAY)
    return p


def callout(doc: Document, text: str, style: str = "note"):
    cfg = {
        "note":    (C_BLUE_LIGHT, C_BLUE_DARK,  "ℹ  참고"),
        "warning": (C_YELLOW_BG,  C_ORANGE,     "⚠  주의"),
        "tip":     (C_GREEN_BG,   C_GREEN_DARK, "✔  핵심"),
    }
    bg, fg, label = cfg.get(style, cfg["note"])
    p = doc.add_paragraph(style="Normal")
    shade_para(p, bg)
    add_left_border(p, color_hex=_rgb_hex(fg), sz=8)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, f"{label}:  ", bold=True, size=10, color=fg)
    add_run(p, text, size=10, color=C_BLACK)
    return p


def image(doc: Document, filename: str, caption: str = None,
          width: float = 5.8, assets_dir: Path = None):
    if assets_dir is None:
        assets_dir = ASSETS_EXEC
    path = assets_dir / filename
    if not path.exists():
        body(doc, f"[이미지 없음: {filename}]", color=RGBColor(0xC6, 0x28, 0x28))
        return
    doc.add_picture(str(path), width=Inches(width))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap_p = doc.add_paragraph(style="Normal")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(2)
        cap_p.paragraph_format.space_after  = Pt(10)
        add_run(cap_p, f"▲  {caption}", italic=True, size=9, color=C_CAPTION)


def image_drawio(doc: Document, filename: str, caption: str = None,
                 width: float = 6.2, style: str = "orig"):
    """draw.io PNG 삽입 — assets/ 폴더에서 로드"""
    assets_dir = ASSETS_NEW
    path = assets_dir / filename
    if not path.exists():
        body(doc, f"[draw.io 이미지 없음: {filename} (style={style})]",
             color=RGBColor(0xC6, 0x28, 0x28))
        return
    doc.add_picture(str(path), width=Inches(width))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap_p = doc.add_paragraph(style="Normal")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(2)
        cap_p.paragraph_format.space_after  = Pt(10)
        add_run(cap_p, f"▲  {caption}", italic=True, size=9, color=C_CAPTION)


def make_table(doc: Document, headers: list, rows: list,
               col_widths_cm: list = None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, C_INDIGO_DARK)
        set_cell_valign_center(cell)
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, color=C_WHITE, size=9.5)

    # 데이터 행
    for ri, row_data in enumerate(rows, 1):
        bg = C_ROW_ALT if ri % 2 == 0 else None
        for ci, val in enumerate(row_data):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            if bg:
                shade_cell(cell, bg)
            set_cell_valign_center(cell)
            p = cell.paragraphs[0]
            add_run(p, str(val), size=9.5)

    # 열 너비
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t


def divider(doc: Document):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    add_bottom_border(p, color_hex="CCCCCC", sz=4)


# ── 목차 (Word 자동 목차 필드) ────────────────────────────────────────────────
def section_toc(doc: Document):
    # h1() 대신 직접 단락 추가 — 목차 제목 자체는 TOC에 포함되지 않도록
    p_title = doc.add_paragraph(style="Normal")
    shade_para(p_title, RGBColor(0xE8, 0xEA, 0xF6))
    add_left_border(p_title, color_hex="0D47A1", sz=20)
    p_title.paragraph_format.left_indent  = Cm(0.4)
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after  = Pt(12)
    add_run(p_title, "목차", bold=True, size=15, color=C_TITLE_BG)

    # ── Word TOC 필드 삽입 ────────────────────────────────────────────────
    # OOXML 표준: fldChar(begin) / instrText / fldChar(separate) / fldChar(end)
    # 각각 별도 <w:r> 에 담아야 Word가 필드로 인식한다.
    # Heading 1~2 기준 자동 목차. Word에서 열면 "필드 업데이트" 프롬프트가 나오면 [업데이트] 클릭.
    # 또는 Ctrl+A → F9 로 수동 업데이트 가능.
    p_toc = doc.add_paragraph(style="Normal")
    p_toc.paragraph_format.space_before = Pt(4)
    p_toc.paragraph_format.space_after  = Pt(4)

    def _add_fld_char(para, type_: str):
        """fldChar 전용 run — 각각 독립 <w:r> 필요"""
        r = para.add_run()
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), type_)
        r._r.append(fc)

    def _add_instr(para, text: str):
        """instrText 전용 run"""
        r = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        # \o "1-2" : Heading 1~2 포함  \h : 하이퍼링크  \z : 웹 레이아웃 숨김  \u : 스타일 사용
        instr.text = text
        r._r.append(instr)

    _add_fld_char(p_toc, "begin")
    _add_instr(p_toc, ' TOC \\o "1-2" \\h \\z \\u ')
    _add_fld_char(p_toc, "separate")
    _add_fld_char(p_toc, "end")

    doc.add_paragraph()   # 목차 아래 여백
    doc.add_page_break()


# ── 변경 이력 ─────────────────────────────────────────────────────────────────
def section_changelog(doc: Document):
    h1(doc, "변경 이력")
    make_table(doc,
        ["날짜", "주요 변경 내용"],
        [
            ["2026-03-19",
             "말풍선(callout) 다이어그램 제거, 표지 레이아웃 개선(제목·이름 간격),\n"
             "2페이지 목차 추가, 3페이지부터 본문 시작으로 구조 재편"],
            ["2026-03-15",
             "단순화(v1.4): ExecutionQueue lock-free 백엔드 제거, BoundedLockFreeQueue 삭제,\n"
             "mutex 단일 백엔드로 통일 (533→157줄). draw.io 다이어그램 전면 갱신,\n"
             "03/04 시퀀스 다이어그램 draw.io 소스 추가."],
            ["2026-03-15",
             "코드 정리(v1.3): 미구현 DBTaskType enum 제거(SaveGameProgress·Custom),\n"
             "미사용 include 제거(<iostream>·<cstring>),\n"
             "Shutdown() 불필요 mutex 제거, 이중 언어 주석 영어 단일화(119파일 −2,180줄)"],
            ["2026-03-15",
             "에러 처리 3-플랫폼 통합(ProcessErrorCompletion), SendResult::InvalidArgument 추가,\n"
             "방향별 에러 통계(totalSendErrors / totalRecvErrors) 분리,\n"
             "DB 워커 수 config화(-w CLI, DEFAULT_DB_WORKER_COUNT / DEFAULT_TASK_QUEUE_WORKER_COUNT),\n"
             "DBTaskQueue.cpp 포맷 정리, Session::Reset() assert, SessionManager 이중 close 방지"],
            ["2026-03-10",
             "섹션 11 신규 추가 — 동기화·락·비동기 핵심 메커니즘.\n"
             "Mermaid 다이어그램 4개 추가."],
            ["2026-03-02",
             "비동기 로직 고도화 A~E, Linux Docker 통합 테스트, AsyncScope 풀 재사용 버그 수정"],
            ["2026-03-01",
             "Core/Memory 버퍼 모듈, 핑퐁 검증, SessionFactory 제거, KeyedDispatcher 도입"],
            ["2026-02-28", "RIO slab pool (WSA 10055 수정), MAX_CONNECTIONS=1000"],
            ["2026-02-26", "초기 작성"],
        ],
        col_widths_cm=[3.0, 15.0]
    )
    divider(doc)
    doc.add_page_break()


# ── 섹션 1: 개요 ──────────────────────────────────────────────────────────────
def section_overview(doc: Document):
    h1(doc, "1. 개요")
    body(doc, (
        "이 보고서는 NetworkModuleTest 프로젝트의 ServerEngine, TestServer, DBServer 구현을 기반으로 "
        "네트워크 처리 구조, 비동기 처리 구조, DB 처리 구조를 코드 수준에서 정리한다. "
        "아키텍처의 각 계층이 어떻게 협력하고, 데이터가 어떤 경로로 흐르며, "
        "비동기·논블로킹 설계가 어떻게 적용되어 있는지를 중심으로 기술한다."
    ))

    image_drawio(doc, "diag_overview.png",
                 caption="전체 아키텍처 개요 — INetworkEngine · AsyncIOProvider · 스레드 역할 분리",
                 style=DRAWIO_STYLE)
    image_drawio(doc, "diag_arch.png",
                 caption="아키텍처 계층 구조", style=DRAWIO_STYLE)

    h2(doc, "1.1 3계층 설계 요약")
    make_table(doc,
        ["계층", "핵심 컴포넌트", "역할 요약"],
        [
            ["네트워크",    "INetworkEngine\nBaseNetworkEngine\nAsyncIOProvider",
             "플랫폼별 I/O 완료 처리\n세션 생성·관리\n송수신 비동기 제공"],
            ["비동기 처리", "ThreadPool\nKeyedDispatcher\nDBTaskQueue\nOrderedTaskQueue",
             "로직 워커 스레드 풀\n세션 키 친화도 라우팅\n논블로킹 DB 오프로딩"],
            ["DB 처리",    "IDatabase\nSQLiteDatabase / MockDatabase\nServerLatencyManager",
             "교체 가능한 DB 추상화\n로컬 SQLite / 테스트 Mock\nDBServer 지연 시간 기록"],
        ],
        col_widths_cm=[3.0, 5.0, 9.0]
    )

    h2(doc, "1.2 플랫폼별 I/O 백엔드")
    make_table(doc,
        ["플랫폼", "1순위 백엔드", "폴백 백엔드", "비고"],
        [
            ["Windows", "RIO (Registered I/O)", "IOCP", "WSA 10055 방지를 위한 Slab 풀"],
            ["Linux",   "io_uring",              "epoll", "kernel 5.1+ 필요"],
            ["macOS",   "kqueue",                "—",    "단일 백엔드"],
        ],
        col_widths_cm=[2.5, 4.5, 3.5, 7.5]
    )
    divider(doc)


# ── 섹션 2: 네트워크 처리 ────────────────────────────────────────────────────
def section_network(doc: Document):
    h1(doc, "2. 네트워크 처리 구조")
    body(doc, (
        "네트워크 계층은 INetworkEngine 인터페이스와 AsyncIOProvider 플랫폼 추상화의 2계층으로 분리된다. "
        "공통 동작(세션 관리, 이벤트 콜백, 통계)은 BaseNetworkEngine이 담당하고, "
        "플랫폼 특화 동작(소켓 accept, completion 처리)은 각 플랫폼 구현체가 담당한다."
    ))

    h2(doc, "2.1 계층 구조")
    make_table(doc,
        ["컴포넌트", "역할", "주요 코드 포인트"],
        [
            ["INetworkEngine",     "외부 API 인터페이스\nInitialize / Start / Stop\nSendData / CloseConnection",
             "Network/Core/NetworkEngine.h:72"],
            ["BaseNetworkEngine",  "공통 구현\n세션 조회·제거\n이벤트 콜백 등록\n통계 집계",
             "Network/Core/BaseNetworkEngine.cpp:28"],
            ["AsyncIOProvider",    "플랫폼 I/O 추상화\nRecvAsync / SendAsync\nAssociateSocket",
             "Network/Core/AsyncIOProvider.cpp:74"],
            ["Platform Engine",   "소켓 accept 루프\nCompletion 처리\n(Windows / Linux / macOS)",
             "Network/Platforms/WindowsNetworkEngine.cpp:128"],
        ],
        col_widths_cm=[4.0, 6.0, 8.0]
    )

    h2(doc, "2.2 연결 수립 → 세션 생성 흐름")
    body(doc, "플랫폼 AcceptLoop()에서 새 연결을 처리하는 공통 흐름은 아래와 같다.")
    numbered(doc, "accept()로 소켓 수락")
    numbered(doc, "SessionManager::CreateSession()으로 세션 객체 생성")
    numbered(doc, "AsyncIOProvider::AssociateSocket()으로 I/O 백엔드 연동")
    numbered(doc, "로직 스레드풀에서 OnConnected + Connected 이벤트 비동기 실행")
    numbered(doc, "첫 번째 Recv 등록 (수신 루프 시작)")

    image_drawio(doc, "diag_client_lifecycle.png",
                 caption="클라이언트 생명주기 시퀀스 — 접속 수락 → 핸드셰이크 → ping/pong → 종료",
                 style=DRAWIO_STYLE)
    image_drawio(doc, "diag_seq.png",
                 caption="연결 수립 시퀀스", style=DRAWIO_STYLE)

    h2(doc, "2.3 Session 구조")
    body(doc, (
        "Session은 네트워크 계층의 핵심 단위다. 연결 상태, 송수신 큐, AsyncScope를 "
        "하나의 객체로 관리하며, 동기화는 역할별로 분리된 프리미티브가 담당한다."
    ))
    image_drawio(doc, "diag_session_uml.png",
                 caption="Session UML — 주요 멤버 변수 및 동기화 프리미티브 구조",
                 style=DRAWIO_STYLE)

    make_table(doc,
        ["동기화 프리미티브", "보호 대상", "설계 포인트"],
        [
            ["mSendMutex (mutex)",        "mSendQueue · mAsyncProvider",
             "락 내 shared_ptr 스냅샷만 복사 후 즉시 해제\n실제 I/O 호출은 락 외부에서 수행"],
            ["mState (atomic, acq_rel)",  "연결 상태 enum",
             "exchange로 Close() TOCTOU 이중 닫기 방지"],
            ["mSocket (atomic, acq_rel)", "소켓 핸들",
             "Close/Send 간 race 방지"],
            ["mIsSending (atomic CAS)",   "이중 전송 방지",
             "compare_exchange_strong(false→true)\n한 스레드만 PostSend() 진입 보장"],
            ["mSendQueueSize (atomic)",   "큐 크기 fast-path 조회",
             "relaxed read (부정확해도 됨)\nrelease store"],
        ],
        col_widths_cm=[4.5, 4.5, 9.0]
    )

    h2(doc, "2.4 수신 처리와 패킷 재조립")
    body(doc, (
        "수신 완료 이벤트는 ProcessCompletions()에서 감지되어 "
        "BaseNetworkEngine::ProcessRecvCompletion()으로 전달된다. "
        "이후 로직 스레드풀로 넘겨진 Session::ProcessRawRecv()에서 TCP 스트림을 재조립한다."
    ))
    bullet(doc, "PacketHeader(size, id) 기준으로 완전한 패킷이 조립될 때까지 버퍼 누적")
    bullet(doc, "유효하지 않은 크기 또는 오버플로우 탐지 시 세션 즉시 종료")
    bullet(doc, "POSIX 경로: RecvAsync()를 직접 구동 (PostRecv() 미사용)")
    code_ref(doc, "Session.cpp:445", "TCP 스트림 재조립 진입점")
    code_ref(doc, "BaseNetworkEngine.cpp:255", "수신 완료 공통 처리")

    h2(doc, "2.5 송신 처리")
    body(doc, (
        "Session::Send()는 SendResult를 반환해 호출자에게 백프레셔 피드백을 제공한다. "
        "mIsSending CAS와 mSendQueueSize atomic으로 불필요한 락 경쟁과 이중 전송을 방지한다."
    ))
    make_table(doc,
        ["SendResult 값", "의미", "호출자 조치"],
        [
            ["Ok",              "전송 성공 (큐잉 또는 즉시 전송)",  "없음"],
            ["QueueFull",       "큐 백프레셔 임계값 초과",           "잠시 후 재시도 또는 세션 종료"],
            ["NotConnected",    "세션 미연결",                       "재시도 금지"],
            ["InvalidArgument", "null 또는 최대 크기 초과 패킷",     "재시도 금지 — 호출 측 버그"],
        ],
        col_widths_cm=[4.0, 6.5, 7.5]
    )
    callout(doc, (
        "bytesSent <= 0인 송신 완료는 DataSent 이벤트를 발생시키지 않고 "
        "ProcessErrorCompletion()으로 라우팅하여 세션을 정상 종료한다. "
        "(BaseNetworkEngine.cpp:ProcessSendCompletion)"
    ), style="note")
    callout(doc, (
        "이벤트 스레드 일관성: OnDisconnected는 CloseConnection() 경로와 recv 오류 경로 모두 "
        "로직 스레드풀로 전달되어 실행된다. 콜백 호출 스레드가 일관하게 유지된다."
    ), style="note")

    h2(doc, "2.6 에러 처리 통합 (ProcessErrorCompletion)")
    body(doc, (
        "3개 플랫폼(Windows/Linux/macOS) 에러 완료 경로를 단일 헬퍼로 통합했다. "
        "Send/Recv 방향 구분, AsyncScope 경유 disconnect, 방향별 통계 집계를 "
        "모든 플랫폼에서 일관되게 처리한다."
    ))
    make_table(doc,
        ["항목", "내용"],
        [
            ["함수 시그니처",
             "ProcessErrorCompletion(SessionRef, AsyncIOType, OSError)"],
            ["방향별 카운터",
             "AsyncIOType::Send → mTotalSendErrors++\n"
             "AsyncIOType::Recv → mTotalRecvErrors++\n"
             "Statistics::totalErrors = sendErrors + recvErrors"],
            ["예상 종료 판별",
             "WSAECONNRESET / WSAESHUTDOWN / EPIPE / ECONNRESET / osError==0\n"
             "→ Logger::Warn (정상 종료)\n그 외 → Logger::Error (비정상)"],
            ["disconnect 경로",
             "ProcessRecvCompletion(session, 0, nullptr) 경유\n"
             "→ AsyncScope 경유 OnDisconnected — 모든 플랫폼 동일"],
            ["코드 포인트",
             "BaseNetworkEngine.cpp:ProcessErrorCompletion"],
        ],
        col_widths_cm=[4.0, 14.0]
    )
    divider(doc)


# ── 섹션 3: 비동기 처리 ──────────────────────────────────────────────────────
def section_async(doc: Document):
    h1(doc, "3. 비동기 처리 구조")
    body(doc, (
        "비동기 처리는 I/O 완료 스레드(플랫폼 엔진)와 로직 스레드풀(KeyedDispatcher)의 "
        "역할 분리를 기반으로 한다. DB 작업은 별도의 논블로킹 큐(DBTaskQueue, OrderedTaskQueue)로 "
        "오프로딩하여 로직 스레드 지연을 최소화한다."
    ))

    image_drawio(doc, "diag_async_1_dispatch.png",
                 caption="I/O 완료 → 워커 배정 흐름", style=DRAWIO_STYLE)

    h2(doc, "3.1 스레드 역할 분리")
    make_table(doc,
        ["스레드 종류", "주체", "역할"],
        [
            ["I/O 완료 스레드",  "플랫폼 엔진\n(Windows/Linux/macOS)",
             "accept / recv / send 완료 감지\n로직 스레드풀로 패킷 전달"],
            ["로직 워커 스레드", "KeyedDispatcher",
             "패킷 처리 · OnConnected · OnDisconnected\n세션 키 친화도 라우팅 (FIFO 순서 보장)"],
            ["DB 워커 스레드 (TestServer)",
             "DBTaskQueue\n기본 1개, -w 플래그로 설정\n(DEFAULT_TASK_QUEUE_WORKER_COUNT)",
             "논블로킹 DB I/O 실행\nWAL 영속성 보장\nsessionId % workerCount 해시 친화도"],
            ["DB 워커 스레드 (DBServer)",
             "OrderedTaskQueue\n기본 4개, -w 플래그로 설정\n(DEFAULT_DB_WORKER_COUNT)",
             "serverId 단위 순서 보장\nKeyedDispatcher 래핑"],
            ["재연결 스레드",    "TestServer DBReconnectLoop",
             "DB 서버 끊김 시 지수 백오프 재시도\nStop() 신호 시 즉시 종료"],
        ],
        col_widths_cm=[4.5, 4.5, 9.0]
    )

    callout(doc, (
        "KeyedDispatcher는 sessionId % workerCount로 항상 같은 워커에 작업을 배분한다. "
        "워커 큐 FIFO와 결합해 세션별 패킷 처리 순서를 보장하며, "
        "Session::mRecvMutex가 불필요해진다."
    ), style="tip")

    h2(doc, "3.2 DBTaskQueue — 논블로킹 DB 오프로딩 (TestServer)")
    body(doc, (
        "TestServer는 접속/해제 시점을 직접 DB에 기록하지 않고 DBTaskQueue에 enqueue한다. "
        "DB I/O 지연이 로직 워커를 블로킹하지 않는다. "
        "워커 수는 CLI -w 플래그로 설정 가능하며 기본값은 1(DEFAULT_TASK_QUEUE_WORKER_COUNT)이다."
    ))
    bullet(doc, "OnClientConnectionEstablished → RecordConnectTime")
    bullet(doc, "OnClientConnectionClosed → RecordDisconnectTime")
    bullet(doc, "sessionId % workerCount 해시 친화도 — 워커 수와 무관하게 세션별 순서 보장")
    bullet(doc, "1 워커: 가장 단순, 친화도 계산 불필요")
    bullet(doc, "N 워커: 처리량 향상, 세션별 순서는 해시 친화도로 여전히 보장")
    callout(doc,
        "워커 수를 바꿔도 순서는 깨지지 않는다. "
        "재시작 시 worker count 변경이 안전하다.",
        style="tip")
    code_ref(doc, "TestServer.cpp:OnClientConnectionEstablished", "비동기 기록 진입점")
    code_ref(doc, "DBTaskQueue.cpp:EnqueueTask", "작업 enqueue")
    code_ref(doc, "DBTaskQueue.cpp:WorkerThreadFunc", "워커 실행 루프")
    code_ref(doc, "NetworkTypes.h:DEFAULT_TASK_QUEUE_WORKER_COUNT", "기본 워커 수 상수")

    h2(doc, "3.3 KeyedDispatcher — 세션 친화 라우팅")
    body(doc, (
        "KeyedDispatcher는 session_id % workerCount 해시를 이용해 동일 세션의 모든 작업을 "
        "항상 같은 Logic Worker 스레드로 라우팅한다. "
        "이를 통해 세션 단위 FIFO 순서가 락 없이 보장되고, "
        "Session::mRecvMutex를 제거할 수 있다."
    ))
    make_table(doc,
        ["항목", "내용"],
        [
            ["배정 공식",     "worker_idx = session_id % workerCount\n결정론적 — 같은 세션은 항상 같은 워커"],
            ["Dispatch()",   "shared_lock(mWorkersMutex)\n다수 I/O 스레드가 동시에 배정 가능"],
            ["Shutdown()",   "exclusive_lock(mWorkersMutex)\n신규 배정 차단 → 남은 작업 완전 드레인"],
            ["ExecutionQueue<T>", "워커당 1개, FIFO 순서 보장\nPush() / Pop(), BackpressurePolicy (RejectNewest / Block)"],
            ["mRecvMutex 제거", "같은 세션 recv가 항상 같은 워커 → 동시 접근 구조적 불가능"],
        ],
        col_widths_cm=[4.5, 13.5]
    )
    callout(doc, (
        "Dispatch()와 Shutdown()은 mWorkersMutex 공유·독점 잠금 쌍으로 안전하게 공존한다. "
        "Shutdown() 진입 즉시 신규 dispatch가 차단되므로 in-flight 작업 유실 없이 종료된다."
    ), style="tip")
    code_ref(doc, "Concurrency/KeyedDispatcher.h", "Dispatch / Shutdown 구현")
    code_ref(doc, "Concurrency/ExecutionQueue.h", "Push / Pop / BackpressurePolicy")

    image_drawio(doc, "diag_async_2_keyed.png",
                 caption="KeyedDispatcher — 세션 친화 라우팅 (session_id % workerCount)",
                 style=DRAWIO_STYLE)

    h2(doc, "3.4 OrderedTaskQueue — DBServer 키 순서 보장")
    body(doc, (
        "TestDBServer는 serverId 단위 작업 순서 보장을 위해 OrderedTaskQueue를 사용한다. "
        "내부적으로 KeyedDispatcher를 래핑하여 같은 key(serverId)를 항상 같은 워커로 라우팅한다. "
        "워커 수는 CLI -w 플래그로 설정 가능하며 기본값은 4(DEFAULT_DB_WORKER_COUNT)이다."
    ))
    bullet(doc, "facade: OrderedTaskQueue.cpp:Initialize()")
    bullet(doc, "keyed dispatch: OrderedTaskQueue.cpp:EnqueueTask()")
    bullet(doc, "dispatcher 구현: Concurrency/KeyedDispatcher.h")
    code_ref(doc, "NetworkTypes.h:DEFAULT_DB_WORKER_COUNT", "기본 워커 수 상수 (기본값 4)")
    code_ref(doc, "DBServer/main.cpp:-w 옵션", "런타임 워커 수 설정")

    h2(doc, "3.5 DB 서버 재연결 루프 (TestServer, Windows 전용)")
    body(doc, "DB 서버 연결이 끊기면 재연결 스레드가 지수 백오프로 재시도한다.")
    make_table(doc,
        ["케이스", "재시도 간격", "비고"],
        [
            ["일반 연결 오류",            "1s → 2s → 4s → ... (최대 30s)", "지수 백오프"],
            ["WSAECONNREFUSED (10061)", "1초 고정",                        "서버 미시작 시 빠른 재시도"],
            ["Stop() 신호",              "즉시 종료",                       "condition_variable로 깨움"],
        ],
        col_widths_cm=[5.0, 5.5, 7.5]
    )
    code_ref(doc, "TestServer.cpp:583", "재연결 루프 진입점")
    code_ref(doc, "TestServer.cpp:629", "오류 종류별 분기")

    divider(doc)


# ── 섹션 4: DB 처리 ──────────────────────────────────────────────────────────
def section_db(doc: Document):
    h1(doc, "4. DB 처리 구조")
    body(doc, (
        "DB 레이어는 IDatabase / IStatement 인터페이스 기반으로 구현체를 교체 가능하게 설계되어 있다. "
        "TestServer는 로컬 SQLite/Mock DB를 DBTaskQueue로 비동기 접근하고, "
        "TestDBServer는 네트워크로 수신한 요청을 OrderedTaskQueue를 통해 처리한다."
    ))

    image_drawio(doc, "diag_db.png",
                 caption="DB 처리 계층 — WAL · Factory · 비동기 큐", style=DRAWIO_STYLE)

    h2(doc, "4.1 IDatabase 추상화 계층")
    make_table(doc,
        ["구현체", "용도", "코드 포인트"],
        [
            ["MockDatabase",    "단위 테스트 / DB 없는 환경",  "ServerEngine/Database/"],
            ["SQLiteDatabase",  "로컬 파일 DB (TestServer)", "ServerEngine/Database/"],
            ["ODBCDatabase",    "ODBC 지원 DB",              "ServerEngine/Database/"],
            ["OLEDBDatabase",   "Windows OLE DB",            "ServerEngine/Database/"],
        ],
        col_widths_cm=[4.0, 6.0, 8.0]
    )
    code_ref(doc, "IDatabase.h:29", "인터페이스 정의")
    code_ref(doc, "DatabaseFactory.cpp:19", "팩토리 생성 진입점")

    h2(doc, "4.2 TestServer 로컬 DB 경로")
    body(doc, "TestServer::Initialize()에서 설정 값에 따라 DB 구현체를 선택한다.")
    make_table(doc,
        ["설정", "선택된 DB", "비고"],
        [
            ["dbConnectionString 비어 있음", "MockDatabase",  "DB 없이 메모리 로그만"],
            ["dbConnectionString 값 있음",  "SQLiteDatabase", "해당 파일 경로 사용"],
        ],
        col_widths_cm=[5.5, 4.0, 8.5]
    )
    body(doc, "선택된 DB 인스턴스를 DBTaskQueue에 주입하고, 큐에서 아래 테이블 존재를 보장한다.")
    bullet(doc, "SessionConnectLog — 접속 시각 기록")
    bullet(doc, "SessionDisconnectLog — 해제 시각 기록")
    bullet(doc, "PlayerData — 플레이어 정보")
    code_ref(doc, "TestServer.cpp:81", "DB 선택 및 주입")
    code_ref(doc, "DBTaskQueue.cpp:73", "테이블 보장 (CREATE TABLE IF NOT EXISTS)")

    h2(doc, "4.3 TestDBServer DB 처리")
    body(doc, (
        "TestDBServer는 ServerPacketHandler + ServerLatencyManager + OrderedTaskQueue 조합으로 동작한다."
    ))
    make_table(doc,
        ["패킷", "처리 흐름", "코드 포인트"],
        [
            ["ServerPingReq",      "RTT 계산 → RecordLatency (메모리)",
             "ServerPacketHandler.cpp:116"],
            ["DBSavePingTimeReq",  "SavePingTime 실행 → 응답 패킷 전송",
             "ServerPacketHandler.cpp:196"],
        ],
        col_widths_cm=[4.5, 7.5, 6.0]
    )
    callout(doc, (
        "현재 TestDBServer 기본 경로에는 ServerLatencyManager::SetDatabase() 주입 코드가 없다. "
        "ExecuteQuery()는 DB 미주입 시 true를 반환하며 로그만 기록한다. "
        "영구 저장은 DBServer.cpp(실험 경로) 또는 주입 코드 추가로 활성화된다."
    ), style="warning")

    h2(doc, "4.4 운영 경로 vs 실험 경로")
    make_table(doc,
        ["경로", "진입점", "DB 주입 여부", "비고"],
        [
            ["TestDBServer (운영)", "DBServer/main.cpp → TestDBServer",
             "미주입 (log-only)", "현재 기본 실행 경로"],
            ["DBServer.cpp (실험)", "DBServer/src/DBServer.cpp:246",
             "ConnectToDatabase()에서 주입",
             "별도 AsyncIO/Protocol 처리 경로"],
        ],
        col_widths_cm=[4.0, 5.5, 4.0, 4.5]
    )
    code_ref(doc, "TestDBServer.cpp:68", "TestDBServer 초기화")
    code_ref(doc, "ServerLatencyManager.cpp:331", "DB 미주입 시 log-only 분기")
    code_ref(doc, "DBServer.cpp:246, 281", "실험 경로 DB 연결 및 주입")
    divider(doc)


# ── 섹션 5: 엔드투엔드 흐름 & Graceful Shutdown ───────────────────────────────
def section_e2e(doc: Document):
    h1(doc, "5. 엔드투엔드 흐름 & Graceful Shutdown")

    h2(doc, "5.1 Client → TestServer 흐름")
    numbered(doc, "클라이언트 TCP 접속")
    numbered(doc, "플랫폼 엔진 accept → Session 생성 → AsyncIOProvider::AssociateSocket()")
    numbered(doc, "로직 스레드: ClientSession::OnConnected() → DBTaskQueue에 접속 시각 enqueue")
    numbered(doc, "클라이언트 SessionConnectReq 처리 → SessionConnectRes 반환")
    numbered(doc, "주기적 PingReq / PongRes 교환")
    numbered(doc, "클라이언트 종료 → OnDisconnected → DBTaskQueue에 해제 시각 enqueue")

    h2(doc, "5.2 TestServer ↔ TestDBServer 흐름")
    numbered(doc, "TestServer: ConnectToDBServer()로 별도 소켓 연결 (Windows 전용)")
    numbered(doc, "DBPingLoop에서 PKT_ServerPingReq 주기 전송")
    numbered(doc, "TestDBServer: ServerPongRes 응답")
    numbered(doc, "주기적으로 PKT_DBSavePingTimeReq 전송 → TestDBServer 응답")
    numbered(doc, "연결 끊김 감지 → DBReconnectLoop 지수 백오프 재시도")

    h2(doc, "5.3 Graceful Shutdown 순서")
    body(doc, "종료 순서는 DB 안전성을 보장하기 위해 명시적으로 정의되어 있다.")

    h3(doc, "TestServer 종료 순서")
    make_table(doc,
        ["순서", "동작", "코드 포인트"],
        [
            ["1", "DB 재연결 루프 깨움 및 join",             "TestServer.cpp:165"],
            ["2", "연결 중 세션 disconnect 기록 enqueue",    "TestServer.cpp:165"],
            ["3", "DBTaskQueue::Shutdown() — 큐 완전 드레인", "TestServer.cpp:165"],
            ["4", "로컬 DB disconnect",                      "TestServer.cpp:165"],
            ["5", "DB 서버 소켓 disconnect",                 "TestServer.cpp:165"],
            ["6", "클라이언트 엔진 Stop()",                  "TestServer.cpp:165"],
        ],
        col_widths_cm=[1.5, 9.0, 7.5]
    )

    h3(doc, "TestDBServer 종료 순서")
    make_table(doc,
        ["순서", "동작", "코드 포인트"],
        [
            ["1", "네트워크 엔진 Stop()",                 "TestDBServer.cpp:160"],
            ["2", "OrderedTaskQueue Shutdown(남은 작업 처리)", "TestDBServer.cpp:160"],
            ["3", "ServerLatencyManager Shutdown()",      "TestDBServer.cpp:160"],
        ],
        col_widths_cm=[1.5, 9.0, 7.5]
    )

    callout(doc, (
        "DBTaskQueue Shutdown()은 남은 모든 작업이 완료될 때까지 블로킹한다. "
        "WAL의 Pending 항목이 모두 Done으로 처리된 후 종료되므로, "
        "정상 종료 시 WAL 재생(crach recovery)이 발생하지 않는다."
    ), style="tip")
    divider(doc)


# ── 섹션 6: 현재 상태 진단 & 개선 제안 ─────────────────────────────────────
def section_assessment(doc: Document):
    h1(doc, "6. 현재 상태 진단 & 개선 제안")

    h2(doc, "6.1 강점")
    bullet(doc, "네트워크 / 로직 / DB 비동기 경로 분리가 명확하며 역할 간 경계가 코드 레벨에서 강제됨")
    bullet(doc, "플랫폼 폴백 체인(RIO→IOCP, io_uring→epoll)으로 최적 백엔드 자동 선택")
    bullet(doc, "재연결 정책(ECONNREFUSED 분리, 지수 백오프)으로 운영 복원력 확보")
    bullet(doc, "KeyedDispatcher 도입으로 mRecvMutex 제거 — 세션 단위 순서 보장과 경쟁 감소 동시 달성")
    bullet(doc, "ProcessErrorCompletion 단일 헬퍼로 3-플랫폼 에러 처리 통합 — Send/Recv 방향 구분, 통계 분리")
    bullet(doc, "SendResult::InvalidArgument 추가 — 재시도 불가 에러와 재시도 가능 에러(QueueFull) 명시적 구분")
    bullet(doc, "DB 워커 수 config화(-w CLI, DEFAULT_DB_WORKER_COUNT / DEFAULT_TASK_QUEUE_WORKER_COUNT) — 매직넘버 제거")

    h2(doc, "6.2 유의점")
    bullet(doc, "TestServer의 DB 서버 소켓 경로는 Windows 전용 (#ifdef _WIN32)")
    bullet(doc, "TestDBServer 기본 경로에 DB 주입 없음 — 영구 저장은 현재 비활성 (log-only)")
    bullet(doc, "TestDBServer 기본 포트(코드 기본 8001)와 스크립트 포트(8002) 혼동 위험")

    h2(doc, "6.3 개선 제안")
    make_table(doc,
        ["우선순위", "개선 항목", "효과"],
        [
            ["High",   "TestDBServer에 설정 기반 DB 주입 경로 추가\n(main -d 옵션 → ServerLatencyManager::SetDatabase())",
             "영구 DB 저장 활성화\n운영/실험 경로 통합"],
            ["Medium", "TestServer DB 서버 연결 경로를\n플랫폼 공통 AsyncIO로 통합",
             "Linux/macOS 지원 확대 (현재 Windows 전용)"],
            ["Low",    "운영 문서에\nTestDBServer vs DBServer.cpp 경로 명시",
             "신규 팀원 혼동 방지"],
            ["완료",   "DB 워커 수 config화 (-w CLI, 상수 정의)\n(DBServer 기본 4, TestServer 기본 1)",
             "매직넘버 제거\n런타임 설정 가능"],
            ["완료",   "3-플랫폼 에러 처리 통합\n(ProcessErrorCompletion, 방향별 통계)",
             "플랫폼 간 동작 일관성 보장\nSend/Recv 에러 가시성 향상"],
        ],
        col_widths_cm=[2.5, 9.0, 6.5]
    )
    divider(doc)


# ── 섹션 7: 주요 참조 파일 ────────────────────────────────────────────────────
def section_references(doc: Document):
    h1(doc, "7. 주요 참조 파일")

    h2(doc, "네트워크 계층")
    code_ref(doc, "Server/ServerEngine/Network/Core/NetworkEngine.h")
    code_ref(doc, "Server/ServerEngine/Network/Core/BaseNetworkEngine.cpp")
    code_ref(doc, "Server/ServerEngine/Network/Core/AsyncIOProvider.cpp")
    code_ref(doc, "Server/ServerEngine/Network/Core/Session.h")
    code_ref(doc, "Server/ServerEngine/Network/Core/Session.cpp")
    code_ref(doc, "Server/ServerEngine/Network/Platforms/WindowsNetworkEngine.cpp")
    code_ref(doc, "Server/ServerEngine/Network/Platforms/LinuxNetworkEngine.cpp")

    h2(doc, "비동기 / 동시성 계층")
    code_ref(doc, "Server/ServerEngine/Utils/NetworkTypes.h", "DEFAULT_DB_WORKER_COUNT / DEFAULT_TASK_QUEUE_WORKER_COUNT")
    code_ref(doc, "Server/ServerEngine/Utils/ThreadPool.h")
    code_ref(doc, "Server/ServerEngine/Concurrency/KeyedDispatcher.h")
    code_ref(doc, "Server/ServerEngine/Concurrency/AsyncScope.h")
    code_ref(doc, "Server/TestServer/src/DBTaskQueue.cpp")
    code_ref(doc, "Server/DBServer/src/OrderedTaskQueue.cpp")

    h2(doc, "DB 계층")
    code_ref(doc, "Server/ServerEngine/Interfaces/IDatabase.h")
    code_ref(doc, "Server/ServerEngine/Database/DatabaseFactory.cpp")
    code_ref(doc, "Server/DBServer/src/TestDBServer.cpp")
    code_ref(doc, "Server/DBServer/src/ServerPacketHandler.cpp")
    code_ref(doc, "Server/DBServer/src/ServerLatencyManager.cpp")
    code_ref(doc, "Server/DBServer/src/DBServer.cpp", "실험 경로")
    code_ref(doc, "Server/TestServer/src/TestServer.cpp")
    code_ref(doc, "Server/TestServer/src/ClientSession.cpp")


# ── 메인 ─────────────────────────────────────────────────────────────────────
def main():
    global DRAWIO_STYLE
    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else \
               Path(__file__).parent.parent / "Network_Async_DB_Report.docx"
    if len(sys.argv) > 2 and sys.argv[2] in ("orig", "B", "C"):
        DRAWIO_STYLE = sys.argv[2]

    print(f"[Build] 보고서 생성 중: {out_path}  (draw.io style={DRAWIO_STYLE})")
    doc = setup_document()
    build_title_page(doc)          # 1페이지: 표지
    section_toc(doc)               # 2페이지: 목차 (+ 페이지 브레이크)
    section_overview(doc)          # 3페이지~: 본문 시작
    section_network(doc)
    section_async(doc)
    section_db(doc)
    section_e2e(doc)
    section_assessment(doc)
    section_references(doc)

    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    print(f"[Done]  {out_path.name}  ({size_kb} KB)")


if __name__ == "__main__":
    main()
